from __future__ import annotations

import html
import re
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


PUBLIC = Path(
    r"C:\Users\johan\Documents\Codex\2026-06-08"
    r"\i-need-to-create-a-module\frontend\public"
)
EFFECTIVE_DATE = "10 June 2026"
REGISTRATION_NUMBER = "2025/674136/07"
REGISTRATION_LINE = (
    f"Inkolo Connect (Pty) Ltd - Registration number {REGISTRATION_NUMBER}"
)

DOCUMENTS = [
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\Buy and sell\BUY & SELL SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "buy-sell-terms",
        "title": "Buy & Sell Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\Catch a ride\CATCH A RIDE SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "catch-a-ride-terms",
        "title": "Catch a Ride Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\funeral cover\# TERMS AND CONDITIONS OF USE.docx"
        ),
        "slug": "funeral-services-terms",
        "title": "Funeral Services Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\Job search\JOB SEARCH SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "job-search-terms",
        "title": "Job Search Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\My community\MY COMMUNITY  CHURCH TERMS AND CONDITIONS.docx"
        ),
        "slug": "my-community-terms",
        "title": "My Community Church Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\VAS\Vas service ts andcs.docx"
        ),
        "slug": "vas-services-terms",
        "title": "VAS Services Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\Keytcha\KEYTCHA PROPERTIES TERMS AND CONDITIONS.docx"
        ),
        "slug": "keytcha-properties-terms",
        "title": "Keytcha Properties Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\Referral\INKOLO REFERRAL SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "referral-service-terms",
        "title": "Inkolo Referral Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\wallet\WALLET SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "wallet-terms",
        "title": "Wallet Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\KZNCC\KZNCC SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "kzncc-service-terms",
        "title": "KZNCC Service Terms and Conditions",
    },
    {
        "source": Path(
            r"C:\Users\johan\OneDrive\Documents\Duranki ecosystem\Services"
            r"\eduU\EDUU SERVICE TERMS AND CONDITIONS.docx"
        ),
        "slug": "eduu-service-terms",
        "title": "EduU Service Terms and Conditions",
    },
]


def paragraph_tag(text: str, index: int) -> str:
    escaped = html.escape(text)
    if index == 0:
        return f"<h1>{escaped}</h1>"
    if index in {1, 2, 3}:
        return f'<p class="document-meta">{escaped}</p>'
    if text in {"Important Disclaimer", "Suggested Platform Checkbox", "Suggested Short Disclaimer for App Footer"}:
        return f"<h2>{escaped}</h2>"
    if re.match(r"^\d+\.\s+\S", text):
        return f"<h2>{escaped}</h2>"
    if re.match(r"^\d+\.\d+(?:\s+\S.*)?$", text):
        return f"<h3>{escaped}</h3>"
    return f"<p>{escaped}</p>"


def normalized_text(text: str) -> str:
    text = re.sub(
        r"Effective Date:\s*_+",
        f"Effective Date: {EFFECTIVE_DATE}",
        text,
        flags=re.IGNORECASE,
    )
    text = re.sub(
        r"Inkolo Connect \(Pty\) Ltd, registration number\s*_+,",
        f"Inkolo Connect (Pty) Ltd, registration number {REGISTRATION_NUMBER},",
        text,
        flags=re.IGNORECASE,
    )
    return text


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    created = Paragraph(new_element, paragraph._parent)
    created.add_run(text)
    return created


for item in DOCUMENTS:
    source = item["source"]
    html_output = PUBLIC / f'{item["slug"]}.html'
    docx_output = PUBLIC / f'{item["slug"]}.docx'
    document = Document(source)
    registration_found = False
    effective_paragraph: Paragraph | None = None
    for paragraph in document.paragraphs:
        updated = normalized_text(paragraph.text)
        if updated != paragraph.text:
            paragraph.text = updated
        if "Effective Date:" in paragraph.text:
            effective_paragraph = paragraph
        if REGISTRATION_NUMBER in paragraph.text:
            registration_found = True

    if not registration_found:
        anchor = effective_paragraph or document.paragraphs[0]
        insert_paragraph_after(anchor, REGISTRATION_LINE)

    paragraphs = [
        normalized_text(paragraph.text.strip())
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    if not any(REGISTRATION_NUMBER in text for text in paragraphs):
        insert_index = min(4, len(paragraphs))
        paragraphs.insert(insert_index, REGISTRATION_LINE)
    body = "\n".join(
        paragraph_tag(text, index) for index, text in enumerate(paragraphs)
    )

    html_output.write_text(
        f"""<!doctype html>
<html lang="en">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>{html.escape(item["title"])}</title>
  <style>
    :root {{ color-scheme: light; }}
    body {{
      max-width: 850px;
      margin: 0 auto;
      padding: 32px 26px 64px;
      color: #17324a;
      background: #fff;
      font: 15px/1.65 Arial, sans-serif;
    }}
    h1 {{ color: #063f91; font-size: 28px; line-height: 1.2; }}
    h2 {{
      margin-top: 30px;
      padding-bottom: 7px;
      color: #0755ad;
      border-bottom: 2px solid #67c72c;
      font-size: 20px;
    }}
    h3 {{ margin-top: 22px; color: #0a4d91; font-size: 16px; }}
    p {{ margin: 9px 0; white-space: pre-wrap; }}
    .document-meta {{ margin: 2px 0; color: #526b7d; font-weight: 700; }}
  </style>
</head>
<body>
{body}
</body>
</html>
""",
        encoding="utf-8",
    )
    document.save(docx_output)
    print(f"Created {html_output}")
    print(f"Created {docx_output}")
